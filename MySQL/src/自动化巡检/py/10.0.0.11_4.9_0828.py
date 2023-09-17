# 导入库
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


#新建空白文档
doc1 = Document()


###############添加页眉#################
header = doc1.sections[0].header # 获取第一个节的页眉 
paragraphaaa = header.paragraphs[0] # 获取页眉的第一个段落 
paragraphaaa.add_run('数据库巡检') # 添加页面内容
##########################################


###############添加页脚#################
footer = doc1.sections[0].footer # 获取第一个节的页脚 
paragraphaaa = footer.paragraphs[0] # 获取页脚的第一个段落 
paragraphaaa.add_run('页脚') # 添加页脚内容
##########################################


###############读取巡检文本内容#################


###1 数据库基础信息###
###操作系统版本###
f1 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_01_os.log",encoding = "utf-8")
###巡检时间###
f2 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_02_time.log",encoding = "utf-8")
f02 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_02_time.log",encoding = "utf-8")
###数据库版本###
f3 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_03_ver.log",encoding = "utf-8")
###数据库名称###
f4 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_04_db.log",encoding = "utf-8")
###数据库创建时间###
f5 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_05_dbctime.log",encoding = "utf-8")
###数据库运行时间###
f6 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_06_runtime.log",encoding = "utf-8")
###数据库字符集###
f7 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_07_char.log",encoding = "utf-8")
###实例字符集###
f8 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_08_char.log",encoding = "utf-8")
###数据库端口号###
f9 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_09_port.log",encoding = "utf-8")
##########################################



###2 数据库文件信息###
###数据库文件信息###
###表空间信息###
f10 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_01_tab.log",encoding = "utf-8")
###undo信息###
f11 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_02_undo.log",encoding = "utf-8")
###临时表空间信息###
f12 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_03_temp.log",encoding = "utf-8")
###redo信息###
f13 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_04_redo.log",encoding = "utf-8")
###redo大小###
f14 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_05_redosize.log",encoding = "utf-8")
###redo组###
f15 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_06_redogroup.log",encoding = "utf-8")
###redo log file目录###
f16 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_07_redodir.log",encoding = "utf-8")
###redo 事务###
f17 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_08_redotrx.log",encoding = "utf-8")
###binlog信息###
f18 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_09_binlog.log",encoding = "utf-8")
###binlog目录位置###
f19 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_10_binlogdir.log",encoding = "utf-8")
###binlog文件列表###
f20 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_11_binloglist.log",encoding = "utf-8")
###binlog状态###
f21 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_12_binlogstatus.log",encoding = "utf-8")
###binglog文件格式###
f22 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_13_binformat.log",encoding = "utf-8")
###binlog缓存###
f23 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_14_bincache.log",encoding = "utf-8")
###binlog使用###
f24 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_15_binuse.log",encoding = "utf-8")
###binlog过期时间###
f25 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_16_binexp.log",encoding = "utf-8")
###binlog大小###
f26 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_17_binsize.log",encoding = "utf-8")
###realy log信息###
f27 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_18_realy.log",encoding = "utf-8")
###realy log 信息###
f28 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_19_realyfile.log",encoding = "utf-8")
###realy log 信息###
f29 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_20_realy.log",encoding = "utf-8")
###配置文件###
f30 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_21_config.log",encoding = "utf-8")
###pid文件###
f31 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_22_pid.log",encoding = "utf-8")
###socket文件###
f32 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_23_socket.log",encoding = "utf-8")
##########################################



###3 数据库高可用信息###
###用户信息###
f54 = open("D:\\python\\log\\10.0.0.11\\03_db_master\\03_01_slave_io.log",encoding = "utf-8")
f55 = open("D:\\python\\log\\10.0.0.11\\03_db_master\\03_02_slave_sql.log",encoding = "utf-8")
f56 = open("D:\\python\\log\\10.0.0.11\\03_db_master\\03_03_behind.log",encoding = "utf-8")
f57 = open("D:\\python\\log\\10.0.0.11\\03_db_master\\03_04_semi.log",encoding = "utf-8")
##########################################


###4 数据库对象信息###
###用户信息###
f33 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_01_user.log",encoding = "utf-8")
f34 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_02_tablemb.log",encoding = "utf-8")
f35 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_03_indexmb.log",encoding = "utf-8")
f36 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_04_creuser.log",encoding = "utf-8")
f37 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_05_tab.log",encoding = "utf-8")
f38 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_06_view.log",encoding = "utf-8")
f39 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_07_tirgger.log",encoding = "utf-8")
f40 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_08_proc.log",encoding = "utf-8")
f41 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_09_part.log",encoding = "utf-8")
f42 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_10_engine.log",encoding = "utf-8")
##########################################


###5 数据库备份信息###
###用户信息###
f43 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_01_cron.log",encoding = "utf-8")
f44 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_02_baksize.log",encoding = "utf-8")
f45 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_03_bakcheck.log",encoding = "utf-8")
##########################################


###6 数据库性能信息###
###用户信息###
f46 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_01_process.log",encoding = "utf-8")
f47 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_02_conn.log",encoding = "utf-8")
f48 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_03_wait.log",encoding = "utf-8")
f49 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_04_trx.log",encoding = "utf-8")
f50 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_05_slow.log",encoding = "utf-8")
f51 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_06_mem.log",encoding = "utf-8")
f52 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_07_mem.log",encoding = "utf-8")
f53 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_08_cpu.log",encoding = "utf-8")
##########################################



###7 数据库参数信息###
###用户信息###
f58 = open("D:\\python\\log\\10.0.0.11\\07_db_para\\07_01_para.log",encoding = "utf-8")
f59 = open("D:\\python\\log\\10.0.0.11\\07_db_para\\07_02_value.log",encoding = "utf-8")
##########################################



###8 数据库系统信息###
###用户信息###
f60 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_01_dfh.log",encoding = "utf-8")
f61 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_02_dfi.log",encoding = "utf-8")
f62 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_03_mount.log",encoding = "utf-8")
f63 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_04_fstab.log",encoding = "utf-8")
##########################################



###9 数据库日志信息###
###用户信息###
f64 = open("D:\\python\\log\\10.0.0.11\\09_db_log\\09_01_err.log",encoding = "utf-8")
f65 = open("D:\\python\\log\\10.0.0.11\\09_db_log\\09_02_message.log",encoding = "utf-8")
##########################################


###############文档标题#################
#添加标题
#先定义标题是head
head = doc1.add_heading(level=0)

#标题内容
title_run = head.add_run('AutoDBCheck系统MySQL数据库巡检')

#标题内容居中
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

#标题英文字体
title_run.font.name = '宋体'

#标题中文字体
title_run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#标题字体颜色
title_run.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#标题字号
#默认一号字体,调整成20号
#from docx.shared import Pt
title_run.font.size = Pt(20)

#字体加粗
title_run.bold = True
##########################################





###############巡检信息#################
#数据库巡检内容
paragraph0 = doc1.add_paragraph('')
run0 = paragraph0.add_run('巡检信息:')
#字号：默认11号，改成15号
run0.font.size = Pt(15)
#字体加粗
run0.bold = True
#标题英文字体
run0.font.name = '宋体'

#标题中文字体
run0.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#标题字体颜色
run0.font.color.rgb = RGBColor(0, 0, 0)      #黑色


#定义表格熟悉(11行，3列)
t0 = doc1.add_table(rows=2, cols=4, style="Table Grid")

#定义表头
cells0 = t0.rows[0].cells
cells0[0].text='序号'
cells0[1].text= '巡检人'
cells0[2].text= '巡检时间'
cells0[3].text= '巡检版本'
#设置表格宽度
#cells[2].width = Inches(10)


#第一行数据
t0.cell(1,0).text='1'
t0.cell(1,1).text='陈举超'
###t0.cell(1,2).text='20220521'
t0.cell(1,2).text=f02.read()
t0.cell(1,3).text='V1.0'
########################################



################IP信息##################
#数据库巡检内容
paragraph1 = doc1.add_paragraph('')
run1 = paragraph1.add_run('\nIP信息：')
#字号：默认11号，改成1号
run1.font.size = Pt(15)
#字体加粗
run1.bold = True
#标题英文字体
run1.font.name = '宋体'
#标题中文字体
run1.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')


#定义表格熟悉(11行，3列)
t1 = doc1.add_table(rows=2, cols=5, style="Table Grid")

#定义表头
cells1 = t1.rows[0].cells
cells1[0].text='序号'
cells1[1].text= '主节点IP'
cells1[2].text= '从节点IP'
cells1[3].text= 'VIP'
cells1[4].text= '数据库架构'
#设置表格宽度
#cells[2].width = Inches(10)

#第一行数据
t1.cell(1,0).text='1'
#t1.cell(1,1).text=f01.read()
#t1.cell(1,2).text=f02.read()
#t1.cell(1,3).text=f03.read()
#t1.cell(1,4).text=f04.read()

t1.cell(1,1).text='10.0.0.11'
t1.cell(1,2).text='10.0.0.13'
t1.cell(1,3).text='10.0.0.15'
t1.cell(1,4).text='MySQL 主从'
########################################


#########数据库巡检内容#################
paragraph = doc1.add_paragraph('')
run = paragraph.add_run('\n数据库巡检内容:')
#字号：默认11号，改成15号
run.font.size = Pt(15)
#字体加粗
run.bold = True
#标题英文字体
run.font.name = '宋体'
#标题中文字体
run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')


#定义表格熟悉(11行，3列)
t = doc1.add_table(rows=11, cols=3, style="Table Grid")

#定义表头
cells2 = t.rows[0].cells
cells2[0].text='序号'
cells2[1].text= '巡检项'
cells2[2].text= '巡检明细'
#设置表格宽度
cells2[2].width = Inches(10)
cells2[1].width = Inches(3)

#第一行数据
t.cell(1,0).text='1'
t.cell(1,1).text='数据库信息'
t.cell(1,2).text='操作系统版本、IP信息、架构信息、数据库创建时间、数据库启动时间、数据库版本信息、数据库补丁信息、字符集等'
#第二行数据
t.cell(2,0).text='2'
t.cell(2,1).text='数据库文件信息'
t.cell(2,2).text='表空间信息、UNDO文件信息、临时文件信息、REDO LOG文件信息、控制文件信息、BINLOG信息等'
#第三行数据
t.cell(3,0).text='3'
t.cell(3,1).text='数据库高可用信息'
t.cell(3,2).text='主从同步信息、同步方式'
#第四行数据
t.cell(4,0).text='4'
t.cell(4,1).text='数据库对象信息'
t.cell(4,2).text='表信息、分区表信息、视图、触发器、存储过程、数据库大小、用户数据大小等'
#第五行数据
t.cell(5,0).text='5'
t.cell(5,1).text='数据库备份信息'
t.cell(5,2).text='备份计划任务、备份大小、备份文件等'
#第六行数据
t.cell(6,0).text='6'
t.cell(6,1).text='数据库性能信息'
t.cell(6,2).text='会话信息、锁信息、TOP SQL等'
#第七行数据
t.cell(7,0).text='7'
t.cell(7,1).text='数据库参数信息'
t.cell(7,2).text='根据参数基线匹配当前参数设置是否正确'
#第八行数据
t.cell(8,0).text='8'
t.cell(8,1).text='数据库系统信息'
t.cell(8,2).text='磁盘目录信息、NFS信息'
#第九行数据
t.cell(9,0).text='9'
t.cell(9,1).text='数据库日志信息'
t.cell(9,2).text='错误日志、慢日志、系统日志'
#第十行数据
t.cell(10,0).text='10'
t.cell(10,1).text='数据库巡检结果'
t.cell(10,2).text='数据库巡检结果'
########################################



#########数据库巡检明细#################
info = doc1.add_heading(level=1)
info_1 = info.add_run('\n数据库巡检明细如下:')

#标题字号
info_1.font.size = Pt(15)
#标题英文字体
info_1.font.name = '宋体'
#标题中文字体
info_1.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
#标题字体颜色
info_1.font.color.rgb = RGBColor(0, 0, 0)      #黑色
#字体加粗
info_1.bold = True
########################################



###########一:数据库基础信息############
#数据库巡检内容
paragraph2 = doc1.add_heading(level = 1)
run2 = paragraph2.add_run('一：数据库基础信息')

#标题英文字体
run2.font.name = '宋体'

#标题中文字体
run2.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run2.font.size = Pt(15)

#标题字体颜色
run2.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run2.bold = True


#定义表格熟悉(11行，3列)
t2 = doc1.add_table(rows=10, cols=2, style="Table Grid")


###20220522 新增表格字体###
t2.style.font.size=Pt(11)  ###字号
#t2.style.font.color.rgb=RGBColor(255, 0, 0) ###字体颜色
#t2.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER ###居中
##########################

#定义表头
cells2 = t2.rows[0].cells
cells2[0].text='数据库基础信息'
cells2[1].text= '内容'
#设置表格宽度
cells2[1].width = Inches(10)


#第1行数据
t2.cell(1,0).text='操作系统版本'
t2.cell(1,1).text=f1.read()

#第2行数据
t2.cell(2,0).text='巡检开始时间'
t2.cell(2,1).text=f2.read()

#第3行数据
t2.cell(3,0).text='数据库版本'
t2.cell(3,1).text=f3.read()

#第4行数据
t2.cell(4,0).text='数据库名称'
t2.cell(4,1).text=f4.read()

#第5行数据
t2.cell(5,0).text='数据库创建时间'
t2.cell(5,1).text=f5.read()

#第6行数据
t2.cell(6,0).text='数据库运行时间'
t2.cell(6,1).text=f6.read()

#第7行数据
t2.cell(7,0).text='数据库字符集'
t2.cell(7,1).text=f7.read()

#第8行数据
t2.cell(8,0).text='实例字符集'
t2.cell(8,1).text=f8.read()

#第9行数据
t2.cell(9,0).text='端口号'
t2.cell(9,1).text=f9.read()
########################################




###########二:数据库文件信息############
#数据库巡检内容
paragraph3 = doc1.add_heading(level = 1)
run3 = paragraph3.add_run('二：数据库文件信息')

#标题英文字体
run3.font.name = '宋体'

#标题中文字体
run3.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run3.font.size = Pt(15)

#标题字体颜色
run3.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run3.bold = True



######表格3#####
#定义表格熟悉(11行，3列)
t3 = doc1.add_table(rows=24, cols=2, style="Table Grid")

#定义表头
cells3 = t3.rows[0].cells
cells3[0].text='数据库文件信息'
cells3[1].text= '内容'

#设置表格宽度
cells3[1].width = Inches(7)

#第1行数据
t3.cell(1,0).text='表空间数量统计'
t3.cell(1,1).text=f10.read()

#第2行数据
t3.cell(2,0).text='UNDO表空间信息'
t3.cell(2,1).text=f11.read()

#第3行数据
t3.cell(3,0).text='TEMOIRARY表空间信息'
t3.cell(3,1).text=f12.read()

#第4行数据
t3.cell(4,0).text='Redo log写缓存'
t3.cell(4,1).text=f13.read()

#第5行数据
t3.cell(5,0).text='Redo log file大小'
t3.cell(5,1).text=f14.read()

#第6数据
t3.cell(6,0).text='Redo log 组数'
t3.cell(6,1).text=f15.read()

#第7数据
t3.cell(7,0).text='Redo log目录位置'
t3.cell(7,1).text=f16.read()

#第8数据
t3.cell(8,0).text='trx_commit'
t3.cell(8,1).text=f17.read()

#第9数据
t3.cell(9,0).text='Binlog是否启用'
t3.cell(9,1).text=f18.read()

#第10数据
t3.cell(10,0).text='Binlog 路径信息'
t3.cell(10,1).text=f19.read()

#第11数据
t3.cell(11,0).text='Binlog 文件列表'
t3.cell(11,1).text=f20.read()

#第12数据
t3.cell(12,0).text='Binlog Position信息'
t3.cell(12,1).text=f21.read()

#第13数据
t3.cell(13,0).text='Binlog格式'
t3.cell(13,1).text=f22.read()

#第14数据
t3.cell(14,0).text='Binlog 缓存'
t3.cell(14,1).text=f23.read()

#第15数据
t3.cell(15,0).text='Binlog 缓存使用情况'
t3.cell(15,1).text=f24.read()

#第16数据
t3.cell(16,0).text='Binlog保留时间'
t3.cell(16,1).text=f25.read()

#第17数据
t3.cell(17,0).text='Binglog单个大小'
t3.cell(17,1).text=f26.read()

#第18数据
t3.cell(18,0).text='Realy log路径'
t3.cell(18,1).text=f27.read()

#第19数据
t3.cell(19,0).text='Realy log文件'
t3.cell(19,1).text=f28.read()

#第20数据
t3.cell(20,0).text='Realy log存储方式'
t3.cell(20,1).text=f29.read()

#第21数据
t3.cell(21,0).text='参数文件'
t3.cell(21,1).text=f30.read()

#第22数据
t3.cell(22,0).text='Pid文件'
t3.cell(22,1).text=f31.read()

#第23数据
t3.cell(23,0).text='socket文件'
t3.cell(23,1).text=f32.read()
########################################




###########三:数据库高可用信息############
#数据库巡检内容
paragraph7 = doc1.add_heading(level = 1)
run7 = paragraph7.add_run('三：数据库高可用信息')

#标题英文字体
run7.font.name = '宋体'

#标题中文字体
run7.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run7.font.size = Pt(15)

#标题字体颜色
run7.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run7.bold = True



######表格4#####
#定义表格熟悉(5行，3列)
t7 = doc1.add_table(rows=5, cols=2, style="Table Grid")

#定义表头
cells7 = t7.rows[0].cells
cells7[0].text='数据库高可用信息'
cells7[1].text= '内容'

#设置表格宽度
cells7[1].width = Inches(7)

#第1行数据
t7.cell(1,0).text='I/O线程'
t7.cell(1,1).text=f54.read()

#第2行数据
t7.cell(2,0).text='SQL线程'
t7.cell(2,1).text=f55.read()

#第3行数据
t7.cell(3,0).text='同步延时'
t7.cell(3,1).text=f56.read()

#第4行数据
t7.cell(4,0).text='同步方式'
t7.cell(4,1).text=f57.read()
########################################



###########四:数据库对象信息############
#数据库巡检内容
paragraph4 = doc1.add_heading(level = 1)
run4 = paragraph4.add_run('四：数据库对象信息')

#标题英文字体
run4.font.name = '宋体'

#标题中文字体
run4.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run4.font.size = Pt(15)

#标题字体颜色
run4.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run4.bold = True



######表格4#####
#定义表格熟悉(11行，3列)
t4 = doc1.add_table(rows=11, cols=2, style="Table Grid")

#定义表头
cells4 = t4.rows[0].cells
cells4[0].text='数据库对象信息'
cells4[1].text= '内容'

#设置表格宽度
cells4[1].width = Inches(7)

#第1行数据
t4.cell(1,0).text='用户信息'
t4.cell(1,1).text=f33.read()

#第2行数据
t4.cell(2,0).text='数据库总大小(MB)'
t4.cell(2,1).text=f34.read()

#第3行数据
t4.cell(3,0).text='索引总大小(MB)'
t4.cell(3,1).text=f35.read()

#第4行数据
t4.cell(4,0).text='用户创建时间'
t4.cell(4,1).text=f36.read()

#第5行数据
t4.cell(5,0).text='用户表数量'
t4.cell(5,1).text=f37.read()

#第6行数据
t4.cell(6,0).text='用户视图数量'
t4.cell(6,1).text=f38.read()

#第7行数据
t4.cell(7,0).text='用户触发器数量'
t4.cell(7,1).text=f39.read()

#第8行数据
t4.cell(8,0).text='用户存储过程数量'
t4.cell(8,1).text=f40.read()

#第9行数据
t4.cell(9,0).text='用户分区表数量'
t4.cell(9,1).text=f41.read()

#第10行数据
t4.cell(10,0).text='用户存储引擎数量'
t4.cell(10,1).text=f42.read()
########################################




###########五:数据库备份信息############
#数据库巡检内容
paragraph5 = doc1.add_heading(level = 1)
run5 = paragraph5.add_run('五：数据库备份信息')

#标题英文字体
run5.font.name = '宋体'

#标题中文字体
run5.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run5.font.size = Pt(15)

#标题字体颜色
run5.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run5.bold = True



######表格5#####
#定义表格熟悉(11行，3列)
t5 = doc1.add_table(rows=4, cols=2, style="Table Grid")

#定义表头
cells5 = t5.rows[0].cells
cells5[0].text='数据库对象信息'
cells5[1].text= '内容'

#设置表格宽度
cells5[1].width = Inches(7)

#第1行数据
t5.cell(1,0).text='数据库本地备份计划任务'
t5.cell(1,1).text=f43.read()

#第2行数据
t5.cell(2,0).text='数据库备份文件大小'
t5.cell(2,1).text=f44.read()

#第3行数据
t5.cell(3,0).text='数据库备份错误信息'
t5.cell(3,1).text=f45.read()
########################################




###########六:数据库性能信息############
#数据库巡检内容
paragraph6 = doc1.add_heading(level = 1)
run6 = paragraph6.add_run('六：数据库性能信息')

#标题英文字体
run6.font.name = '宋体'

#标题中文字体
run6.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run6.font.size = Pt(15)

#标题字体颜色
run6.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run6.bold = True



######表格5#####
#定义表格熟悉(9行，3列)
t6 = doc1.add_table(rows=9, cols=2, style="Table Grid")

#定义表头
cells6 = t6.rows[0].cells
cells6[0].text='数据库性能信息'
cells6[1].text= '内容'

#设置表格宽度
cells6[1].width = Inches(7)

#第1行数据
t6.cell(1,0).text='process信息'
t6.cell(1,1).text=f46.read()

#第2行数据
t6.cell(2,0).text='最大连接信息'
t6.cell(2,1).text=f47.read()

#第3行数据
t6.cell(3,0).text='等待事件信息'
t6.cell(3,1).text=f48.read()

#第4行数据
t6.cell(4,0).text='事务信息'
t6.cell(4,1).text=f49.read()

#第5行数据
t6.cell(5,0).text='慢日志信息'
t6.cell(5,1).text=f50.read()

#第6行数据
t6.cell(6,0).text='内存使用率'
t6.cell(6,1).text=f51.read()

#第7行数据
t6.cell(7,0).text='内存使用TOP 10'
t6.cell(7,1).text=f52.read()

#第8行数据
t6.cell(8,0).text='CPU使用TOP 10'
t6.cell(8,1).text=f53.read()
########################################



###########七:数据库参数信息############
#数据库巡检内容
paragraph8 = doc1.add_heading(level = 1)
run8 = paragraph8.add_run('七：数据库参数信息')

#标题英文字体
run8.font.name = '宋体'

#标题中文字体
run8.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run8.font.size = Pt(15)

#标题字体颜色
run8.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run8.bold = True



######表格7#####
#定义表格熟悉(5行，3列)
t8 = doc1.add_table(rows=2, cols=2, style="Table Grid")

#定义表头
cells8 = t8.rows[0].cells
cells8[0].text='参数'
cells8[1].text= '值'

#设置表格宽度
cells8[1].width = Inches(7)

#第1行数据
t8.cell(1,0).text=f58.read()
t8.cell(1,1).text=f59.read()

#第2行数据
#t8.cell(2,0).text='值'
#t8.cell(2,1).text=f59.read()
########################################



###########八:数据库系统信息############
#数据库巡检内容
paragraph9 = doc1.add_heading(level = 1)
run9 = paragraph9.add_run('八：数据库系统信息')

#标题英文字体
run9.font.name = '宋体'

#标题中文字体
run9.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run9.font.size = Pt(15)

#标题字体颜色
run9.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run9.bold = True



######表格8#####
#定义表格熟悉(5行，3列)
t9 = doc1.add_table(rows=5, cols=2, style="Table Grid")

#定义表头
cells9 = t9.rows[0].cells
cells9[0].text='数据库系统信息'
cells9[1].text= '内容'

#设置表格宽度
cells9[1].width = Inches(7)

#第1行数据
t9.cell(1,0).text='磁盘空间'
t9.cell(1,1).text=f60.read()

#第2行数据
t9.cell(2,0).text='inode空间'
t9.cell(2,1).text=f61.read()

#第3行数据
t9.cell(3,0).text='mount信息'
t9.cell(3,1).text=f62.read()

#第4行数据
t9.cell(4,0).text='开机自动挂载'
t9.cell(4,1).text=f63.read()
########################################



###########九:数据库日志信息############
#数据库巡检内容
paragraph10 = doc1.add_heading(level = 1)
run10 = paragraph10.add_run('九：数据库日志信息')

#标题英文字体
run10.font.name = '宋体'

#标题中文字体
run10.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run10.font.size = Pt(15)

#标题字体颜色
run10.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run10.bold = True



######表格8#####
#定义表格熟悉(5行，3列)
t10 = doc1.add_table(rows=3, cols=2, style="Table Grid")

#定义表头
cells10 = t10.rows[0].cells
cells10[0].text='数据库日志信息'
cells10[1].text= '内容'

#设置表格宽度
cells10[1].width = Inches(7)

#第1行数据
t10.cell(1,0).text='error日志'
t10.cell(1,1).text=f64.read()

#第2行数据
t10.cell(2,0).text='message日志'
t10.cell(2,1).text=f65.read()
########################################



#保存文档
doc1.save('20220828-AutoDBCheck系统MySQL数据库巡检报告-4.9.docx')



#关闭文件
f1.close()
f2.close()
f02.close()
f3.close()
f4.close()
f5.close()
f6.close()
f7.close()
f8.close()
f9.close()
f10.close()
f11.close()
f12.close()
f13.close()
f14.close()
f15.close()
f16.close()
f17.close()
f18.close()
f19.close()
f20.close()
f21.close()
f22.close()
f23.close()
f24.close()
f25.close()
f26.close()
f27.close()
f28.close()
f29.close()
f30.close()
f31.close()
f32.close()
f33.close()
f34.close()
f35.close()
f36.close()
f37.close()
f38.close()
f39.close()
f40.close()
f41.close()
f42.close()
f43.close()
f44.close()
f45.close()
f46.close()
f47.close()
f48.close()
f49.close()
f50.close()
f51.close()
f52.close()
f53.close()
f54.close()
f55.close()
f56.close()
f57.close()
f58.close()
f59.close()
f60.close()
f61.close()
f62.close()
f63.close()
f64.close()
f65.close()


print ("巡检报告已生成到D:\python\py目录下")




