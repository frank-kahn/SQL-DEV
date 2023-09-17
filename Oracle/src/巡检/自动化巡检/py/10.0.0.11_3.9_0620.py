# 导入库
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


#新建空白文档
doc1 = Document()


###############添加页眉#################
header = doc1.sections[0].header # 获取第一个节的页眉 
paragraphaaa = header.paragraphs[0] # 获取页眉的第一个段落 
paragraphaaa.add_run('数据库巡检') # 添加页面内容
##########################################


###############添加页脚#################
footer = doc1.sections[0].footer # 获取第一个节的页脚 
paragraphaaa = footer.paragraphs[0] # 获取页脚的第一个段落 
paragraphaaa.add_run('页脚') # 添加页脚内容
##########################################


###############读取巡检文本内容#################


###1 数据库基础信息###
###操作系统版本###
f1 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_01_os.log",encoding = "utf-8")
###巡检时间###
f2 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_02_time.log",encoding = "utf-8")
f02 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_02_time.log",encoding = "utf-8")
###数据库版本###
f3 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_03_ver.log",encoding = "utf-8")
###数据库补丁###
f4 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_04_path.log",encoding = "utf-8")
###数据库创建时间###
f5 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_05_cre.log",encoding = "utf-8")
###数据库启动时间###
f6 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_06_start.log",encoding = "utf-8")
###数据库实例状态###
f7 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_07_ins.log",encoding = "utf-8")
###数据库字符集###
f8 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_08_char.log",encoding = "utf-8")
###数据库实例名###
f9 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_09_insname.log",encoding = "utf-8")
###数据库DBID###
f10 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_10_dbid.log",encoding = "utf-8")
###数据库端口信息###
f11 = open("D:\\python\\log\\10.0.0.11\\01_db_info\\01_11_port.log",encoding = "utf-8")
##########################################



###2 数据库文件信息###
###数据库文件信息###
###表空间名称###
f12 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_01_name.log",encoding = "utf-8")
###表空间状态###
f13 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_02_tabstat.log",encoding = "utf-8")
###表空间总大小###
f14 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_03_total.log",encoding = "utf-8")
###表空间使用率###
f15 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_04_use.log",encoding = "utf-8")
###数据文件###
f16 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_05_file.log",encoding = "utf-8")
###自动扩展###
f17 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_06_auto.log",encoding = "utf-8")
###扩展最大值###
f18 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_07_max.log",encoding = "utf-8")
###数据文件状态###
f19 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_08_state.log",encoding = "utf-8")
###数据文件名称###
f20 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_09_name.log",encoding = "utf-8")
###回滚段###
f21 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_10_roll.log",encoding = "utf-8")
###临时文件###
f22 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_11_temp.log",encoding = "utf-8")
###控制文件###
f23 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_12_con.log",encoding = "utf-8")
###redo log file###
f24 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_13_redo.log",encoding = "utf-8")
###redo###
f25 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_14_redo.log",encoding = "utf-8")
###redo切换频率###
f26 = open("D:\\python\\log\\10.0.0.11\\02_db_file\\02_15_swi.log",encoding = "utf-8")
##########################################



###3 数据库RAC信息###
###ASM磁盘组###
f27 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_01_asm.log",encoding = "utf-8")
###ASM磁盘组###
f28 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_02_asmdisk.log",encoding = "utf-8")
###ASM磁盘组###
f29 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_03_ocr.log",encoding = "utf-8")
###ASM磁盘组###
f30 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_04_ocrcheck.log",encoding = "utf-8")
###ASM磁盘组###
f31 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_05_ocrbak.log",encoding = "utf-8")
###ASM磁盘组###
f32 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_06_vote.log",encoding = "utf-8")
###ASM磁盘组###
f33 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_07_clu.log",encoding = "utf-8")
###ASM磁盘组###
f34 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_08_res.log",encoding = "utf-8")
###ASM磁盘组###
f35 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_09_scan.log",encoding = "utf-8")
###ASM磁盘组###
f36 = open("D:\\python\\log\\10.0.0.11\\03_db_rac\\03_10_lis.log",encoding = "utf-8")
##########################################



###4 数据库对象信息###
###异常对象信息###
f37 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_01_ao.log",encoding = "utf-8")
###拥有DBA角色的用户###
f38 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_02_role.log",encoding = "utf-8")
###密码策略###
f39 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_03_pass.log",encoding = "utf-8")
###数据库总大小###
f40 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_04_size.log",encoding = "utf-8")
###用户数据大小###
f41 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_05_usize.log",encoding = "utf-8")
###用户名###
f42 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_06_username.log",encoding = "utf-8")
###用户创建时间###
f43 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_07_usertime.log",encoding = "utf-8")
###用户默认表空间###
f44 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_08_usertabo.log",encoding = "utf-8")
###用户状态信息###
f45 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_09_userstat.log",encoding = "utf-8")
###用户对象汇总###
f46 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_10_object.log",encoding = "utf-8")
###失效对象###
f47 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_11_inv.log",encoding = "utf-8")
###分区表###
f48 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_12_par.log",encoding = "utf-8")
###job###
f49 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_13_job.log",encoding = "utf-8")
###SCHEDULER_JOBS###
f50 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_14_sch.log",encoding = "utf-8")
###dblinks###
f51 = open("D:\\python\\log\\10.0.0.11\\04_db_object\\04_15_link.log",encoding = "utf-8")
##########################################


###5 数据库备份信息###
###全备信息###
f52 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_01_fullbak.log",encoding = "utf-8")
###全备信息###
f53 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_02_levelbak.log",encoding = "utf-8")
###全备信息###
f54 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_03_archbak.log",encoding = "utf-8")
###全备信息###
f55 = open("D:\\python\\log\\10.0.0.11\\05_db_bakchk\\05_04_rman.log",encoding = "utf-8")
##########################################



###6 数据库性能信息###
###会话信息###
f56 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_01_session.log",encoding = "utf-8")
###会话信息###
f57 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_02_res.log",encoding = "utf-8")
###资源信息###
f58 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_03_task.log",encoding = "utf-8")
###sga等信息###
f59 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_04_sga.log",encoding = "utf-8")
###等待事件信息###
f60 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_05_event.log",encoding = "utf-8")
###内存信息###
f61 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_06_free.log",encoding = "utf-8")
###内存前十信息###
f62 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_07_memuse.log",encoding = "utf-8")
###CPU前十信息###
f63 = open("D:\\python\\log\\10.0.0.11\\06_db_per\\06_08_cpuuse.log",encoding = "utf-8")
##########################################


###7 数据库参数信息###
###参数信息#
f64 = open("D:\\python\\log\\10.0.0.11\\07_db_para\\07_01_para.log",encoding = "utf-8")
##########################################


###8 数据库操作系统信息###
###磁盘信息###
f65 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_01_df.log",encoding = "utf-8")
###挂载信息###
f66 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_02_mount.log",encoding = "utf-8")
###自动挂载配置###
f67 = open("D:\\python\\log\\10.0.0.11\\08_db_os\\08_03_fstab.log",encoding = "utf-8")
##########################################


###9 数据库日志信息###
###告警日志信息###
f68 = open("D:\\python\\log\\10.0.0.11\\09_db_log\\09_01_log.log",encoding = "utf-8")
###message日志###
f69 = open("D:\\python\\log\\10.0.0.11\\09_db_log\\09_02_message.log",encoding = "utf-8")
##########################################


###############文档标题#################
#添加标题
#先定义标题是head
head = doc1.add_heading(level=0)

#标题内容
title_run = head.add_run('AutoDBCheck系统数据库巡检')

#标题内容居中
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

#标题英文字体
title_run.font.name = '宋体'

#标题中文字体
title_run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#标题字体颜色
title_run.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#标题字号
#默认一号字体,调整成20号
#from docx.shared import Pt
title_run.font.size = Pt(20)

#字体加粗
title_run.bold = True
##########################################





###############巡检信息#################
#数据库巡检内容
paragraph0 = doc1.add_paragraph('')
run0 = paragraph0.add_run('巡检信息:')
#字号：默认11号，改成15号
run0.font.size = Pt(15)
#字体加粗
run0.bold = True
#标题英文字体
run0.font.name = '宋体'

#标题中文字体
run0.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#标题字体颜色
run0.font.color.rgb = RGBColor(0, 0, 0)      #黑色


#定义表格熟悉(11行，3列)
t0 = doc1.add_table(rows=2, cols=4, style="Table Grid")

#定义表头
cells0 = t0.rows[0].cells
cells0[0].text='序号'
cells0[1].text= '巡检人'
cells0[2].text= '巡检时间'
cells0[3].text= '巡检版本'
#设置表格宽度
#cells[2].width = Inches(10)


#第一行数据
t0.cell(1,0).text='1'
t0.cell(1,1).text='陈举超'
###t0.cell(1,2).text='20220521'
t0.cell(1,2).text=f02.read()
t0.cell(1,3).text='V1.0'
########################################



################IP信息##################
#数据库巡检内容
paragraph1 = doc1.add_paragraph('')
run1 = paragraph1.add_run('\nIP信息：')
#字号：默认11号，改成1号
run1.font.size = Pt(15)
#字体加粗
run1.bold = True
#标题英文字体
run1.font.name = '宋体'
#标题中文字体
run1.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')


#定义表格熟悉(11行，3列)
t1 = doc1.add_table(rows=2, cols=5, style="Table Grid")

#定义表头
cells1 = t1.rows[0].cells
cells1[0].text='序号'
cells1[1].text= '节点1_IP'
cells1[2].text= '节点2_IP'
cells1[3].text= 'SCAN IP'
cells1[4].text= '数据库架构'
#设置表格宽度
#cells[2].width = Inches(10)

#第一行数据
t1.cell(1,0).text='1'
#t1.cell(1,1).text=f01.read()
#t1.cell(1,2).text=f02.read()
#t1.cell(1,3).text=f03.read()
#t1.cell(1,4).text=f04.read()

t1.cell(1,1).text='10.0.0.11'
t1.cell(1,2).text='10.0.0.13'
t1.cell(1,3).text='10.0.0.15'
t1.cell(1,4).text='Oracle RAC'
########################################


#########数据库巡检内容#################
paragraph = doc1.add_paragraph('')
run = paragraph.add_run('\n数据库巡检内容:')
#字号：默认11号，改成15号
run.font.size = Pt(15)
#字体加粗
run.bold = True
#标题英文字体
run.font.name = '宋体'
#标题中文字体
run.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')


#定义表格熟悉(11行，3列)
t = doc1.add_table(rows=11, cols=3, style="Table Grid")

#定义表头
cells2 = t.rows[0].cells
cells2[0].text='序号'
cells2[1].text= '巡检项'
cells2[2].text= '巡检明细'
#设置表格宽度
cells2[2].width = Inches(10)
cells2[1].width = Inches(3)

#第一行数据
t.cell(1,0).text='1'
t.cell(1,1).text='数据库信息'
t.cell(1,2).text='操作系统版本、IP信息、VIP信息、架构信息、数据库创建时间、数据库启动时间、数据库版本信息、数据库补丁信息、字符集、实例状态、实例名、DBID、端口号'
#第二行数据
t.cell(2,0).text='2'
t.cell(2,1).text='数据库文件信息'
t.cell(2,2).text='表空间信息、数据文件信息、临时文件信息、UNDO文件信息、回滚段信息、REDO LOG文件信息、ARCHIVELOG信息、控制文件信息'
#第三行数据
t.cell(3,0).text='3'
t.cell(3,1).text='数据库高可用信息'
t.cell(3,2).text='VCS进程信息、VCS配置信息、VCS资源信息、VCS冻结信息、ASM磁盘组信息、ASM磁盘信息、OCR信息、OLR信息、VOTEDISK信息、CRS信息、监听信息'
#第四行数据
t.cell(4,0).text='4'
t.cell(4,1).text='数据库对象信息'
t.cell(4,2).text='表信息、分区表信息、视图、触发器、存储过程、索引、失效对象、异常对象、JOBS、SCHEDULER_JOBS、数据库大小、用户数据大小'
#第五行数据
t.cell(5,0).text='5'
t.cell(5,1).text='数据库备份信息'
t.cell(5,2).text='全备信息、增量备份信息、归档备份信息、异地备份文件信息'
#第六行数据
t.cell(6,0).text='6'
t.cell(6,1).text='数据库性能信息'
t.cell(6,2).text='用户角色、用户密码策略、用户统计信息、等待事件、TOP SQL'
#第七行数据
t.cell(7,0).text='7'
t.cell(7,1).text='数据库参数信息'
t.cell(7,2).text='根据参数基线匹配当前参数设置是否正确'
#第八行数据
t.cell(8,0).text='8'
t.cell(8,1).text='数据库系统信息'
t.cell(8,2).text='磁盘目录信息、NFS信息'
#第九行数据
t.cell(9,0).text='9'
t.cell(9,1).text='数据库日志信息'
t.cell(9,2).text='数据库告警日志、集群告警日志、ASM日志、操作系统日志'
#第十行数据
t.cell(10,0).text='10'
t.cell(10,1).text='数据库巡检结果'
t.cell(10,2).text='数据库巡检结果'
########################################



#########数据库巡检明细#################
info = doc1.add_heading(level=1)
info_1 = info.add_run('\n数据库巡检明细如下:')

#标题字号
info_1.font.size = Pt(15)
#标题英文字体
info_1.font.name = '宋体'
#标题中文字体
info_1.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
#标题字体颜色
info_1.font.color.rgb = RGBColor(0, 0, 0)      #黑色
#字体加粗
info_1.bold = True
########################################



###########一:数据库基础信息############
#数据库巡检内容
paragraph2 = doc1.add_heading(level = 1)
run2 = paragraph2.add_run('一：数据库基础信息')

#标题英文字体
run2.font.name = '宋体'

#标题中文字体
run2.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run2.font.size = Pt(15)

#标题字体颜色
run2.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run2.bold = True


#定义表格熟悉(11行，3列)
t2 = doc1.add_table(rows=12, cols=2, style="Table Grid")


###20220522 新增表格字体###
t2.style.font.size=Pt(11)  ###字号
#t2.style.font.color.rgb=RGBColor(255, 0, 0) ###字体颜色
#t2.style.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER ###居中
##########################

#定义表头
cells2 = t2.rows[0].cells
cells2[0].text='数据库基础信息'
cells2[1].text= '内容'
#设置表格宽度
cells2[1].width = Inches(10)


#第1行数据
t2.cell(1,0).text='操作系统版本'
t2.cell(1,1).text=f1.read()

#第2行数据
t2.cell(2,0).text='巡检开始时间'
t2.cell(2,1).text=f2.read()

#第3行数据
t2.cell(3,0).text='数据库版本'
t2.cell(3,1).text=f3.read()

#第4行数据
t2.cell(4,0).text='数据库补丁信息'
t2.cell(4,1).text=f4.read()

#第5行数据
t2.cell(5,0).text='数据库创建时间'
t2.cell(5,1).text=f5.read()

#第6行数据
t2.cell(6,0).text='数据库启动时间'
t2.cell(6,1).text=f6.read()

#第7行数据
t2.cell(7,0).text='实例状态'
t2.cell(7,1).text=f7.read()

#第8行数据
t2.cell(8,0).text='字符集'
t2.cell(8,1).text=f8.read()

#第9行数据
t2.cell(9,0).text='实例名'
t2.cell(9,1).text=f9.read()

#第10行数据
t2.cell(10,0).text='DBID'
t2.cell(10,1).text=f10.read()

#第11行数据
t2.cell(11,0).text='端口号'
t2.cell(11,1).text=f11.read()
########################################




###########二:数据库文件信息############
#数据库巡检内容
paragraph3 = doc1.add_heading(level = 1)
run3 = paragraph3.add_run('二：数据库文件信息')

#标题英文字体
run3.font.name = '宋体'

#标题中文字体
run3.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run3.font.size = Pt(15)

#标题字体颜色
run3.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run3.bold = True



######表格1#####
#定义表格熟悉(11行，3列)
t3 = doc1.add_table(rows=2, cols=4, style="Table Grid")

#定义表头
cells3 = t3.rows[0].cells
#cells3(0,0).text='表空间名称'
#cells3[0,1].text= '状态'
#cells3[0,2].text= '大小'
#cells3[0,3].text= '使用率'

#设置表格宽度
cells3[1].width = Inches(7)

#第一行数据
t3.cell(0,0).text='表空间名称'
t3.cell(0,1).text='状态'
t3.cell(0,2).text='大小(MB)'
t3.cell(0,3).text='表空间使用率%'

#第二行数据
t3.cell(1,0).text=f12.read()
t3.cell(1,1).text=f13.read()
t3.cell(1,2).text=f14.read()
t3.cell(1,3).text=f15.read()



######表格2#####
#定义表格熟悉(11行，3列)
t4 = doc1.add_table(rows=2, cols=4, style="Table Grid")

#定义表头
cells4 = t4.rows[0].cells

#设置表格宽度
#cells4[1].width = Inches(7)

#第一行数据
t4.cell(0,0).text='数据库文件大小'
#t4.cell(0,1).text='可扩展'
t4.cell(0,1).text='最大可扩展'
t4.cell(0,2).text='数据文件状态'
t4.cell(0,3).text='数据文件路径'

#第二行数据
t4.cell(1,0).text=f16.read()
t4.cell(1,1).text=f18.read()
t4.cell(1,2).text=f19.read()
t4.cell(1,3).text=f20.read()

######表格3#####
#定义表格熟悉(11行，3列)
t5 = doc1.add_table(rows=7, cols=2, style="Table Grid")

#定义表头
cells5 = t5.rows[0].cells
cells5[0].text='数据库文件信息'
cells5[1].text= '内容'

#设置表格宽度
cells5[1].width = Inches(7)

#第二行数据
t5.cell(1,0).text='回滚段名称'
t5.cell(1,1).text=f21.read()

#第三行数据
t5.cell(2,0).text='临时文件'
t5.cell(2,1).text=f22.read()

#第四行数据
t5.cell(3,0).text='控制文件'
t5.cell(3,1).text=f23.read()

#第五行数据
t5.cell(4,0).text='redo名称'
t5.cell(4,1).text=f24.read()

#第六行数据
t5.cell(5,0).text='redo信息'
t5.cell(5,1).text=f25.read()

#第七数据
t5.cell(6,0).text='日志切换'
t5.cell(6,1).text=f26.read()
########################################


###########三:数据库高可用信息############
#数据库巡检内容
paragraph5 = doc1.add_heading(level = 1)
run5 = paragraph5.add_run('三：数据库高可用信息')

#标题英文字体
run5.font.name = '宋体'

#标题中文字体
run5.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run5.font.size = Pt(15)

#标题字体颜色
run5.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run5.bold = True



#定义表格熟悉(11行，3列)
t5 = doc1.add_table(rows=11, cols=2, style="Table Grid")

#定义表头
cells5 = t5.rows[0].cells
cells5[0].text='数据库高可用信息'
cells5[1].text= '内容'

#设置表格宽度
cells5[1].width = Inches(10)


#第一行数据
t5.cell(1,0).text='ASM磁盘组使用率'
t5.cell(1,1).text=f27.read()
#第二行数据
t5.cell(2,0).text='ASM磁盘状态'
t5.cell(2,1).text=f28.read()
#第三行数据
t5.cell(3,0).text='OCR配置信息'
t5.cell(3,1).text=f29.read()
#第四行数据
t5.cell(4,0).text='OCR状态'
t5.cell(4,1).text=f30.read()
#第五行数据
t5.cell(5,0).text='OCR备份'
t5.cell(5,1).text=f31.read()
#第六行数据
t5.cell(6,0).text='VOTEDISK信息'
t5.cell(6,1).text=f32.read()
#第七行数据
t5.cell(7,0).text='集群信息'
t5.cell(7,1).text=f33.read()
#第八行数据
t5.cell(8,0).text='资源信息'
t5.cell(8,1).text=f34.read()
#第九行数据
t5.cell(9,0).text='sacn信息'
t5.cell(9,1).text=f35.read()
#第十行数据
t5.cell(10,0).text='监听信息'
t5.cell(10,1).text=f36.read()
########################################


###########四:数据库对象信息############
#数据库巡检内容
paragraph6 = doc1.add_heading(level = 1)
run6 = paragraph6.add_run('四：数据库对象信息')

#标题英文字体
run6.font.name = '宋体'

#标题中文字体
run6.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run6.font.size = Pt(15)

#标题字体颜色
run6.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run6.bold = True



#定义表格熟悉(11行，3列)
t6 = doc1.add_table(rows=15, cols=2, style="Table Grid")

#定义表头
cells6 = t6.rows[0].cells
cells6[0].text='数据库对象信息'
cells6[1].text= '内容'

#设置表格宽度
cells6[1].width = Inches(10)


#第一行数据
t6.cell(1,0).text='异常对象'
t6.cell(1,1).text=f37.read()

#第二行数据
t6.cell(2,0).text='拥有DBA角色的用户'
t6.cell(2,1).text=f38.read()

#第三行数据
t6.cell(3,0).text='密码策略'
t6.cell(3,1).text=f39.read()

#第四行数据
t6.cell(4,0).text='数据库总大小(MB)'
t6.cell(4,1).text=f40.read()

#第五行数据
t6.cell(5,0).text='用户总大小(MB)'
t6.cell(5,1).text=f41.read()

#第六行数据
t6.cell(6,0).text='用户名'
t6.cell(6,1).text=f42.read()

#第七行数据
t6.cell(7,0).text='用户创建时间'
t6.cell(7,1).text=f43.read()

#第八行数据
t6.cell(8,0).text='用户默认表空间'
t6.cell(8,1).text=f44.read()

#第九行数据
t6.cell(9,0).text='用户状态'
t6.cell(9,1).text=f45.read()

#第十行数据
t6.cell(10,0).text='用户对象数量'
t6.cell(10,1).text=f46.read()

#第十一行数据
t6.cell(11,0).text='失效对象'
t6.cell(11,1).text=f47.read()

#第十二行数据
t6.cell(12,0).text='分区表'
t6.cell(12,1).text=f48.read()

#第十三行数据
t6.cell(13,0).text='job信息'
t6.cell(13,1).text=f49.read()


#第十四行数据
t6.cell(14,0).text='scheduler jobs'
t6.cell(14,1).text=f50.read()

#第十五行数据
t6.cell(14,0).text='dblinks'
t6.cell(14,1).text=f51.read()
########################################



###########五:数据库备份信息############
#数据库巡检内容
paragraph7 = doc1.add_heading(level = 1)
run7 = paragraph7.add_run('五：数据库备份信息')

#标题英文字体
run7.font.name = '宋体'

#标题中文字体
run7.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run7.font.size = Pt(15)

#标题字体颜色
run7.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run7.bold = True



#定义表格熟悉(11行，3列)
t7 = doc1.add_table(rows=5, cols=2, style="Table Grid")

#定义表头
cells7 = t7.rows[0].cells
cells7[0].text='数据库备份信息'
cells7[1].text= '内容'

#设置表格宽度
cells7[1].width = Inches(10)


#第一行数据
t7.cell(1,0).text='全备信息'
t7.cell(1,1).text=f52.read()

#第二行数据
t7.cell(2,0).text='增量备份信息'
t7.cell(2,1).text=f53.read()

#第三行数据
t7.cell(3,0).text='归档备份信息'
t7.cell(3,1).text=f54.read()

#第四行数据
t7.cell(4,0).text='rman配置信息'
t7.cell(4,1).text=f55.read()
########################################


###########六:数据库性能信息############
#数据库巡检内容
paragraph8 = doc1.add_heading(level = 1)
run8 = paragraph8.add_run('六：数据库性能信息')

#标题英文字体
run8.font.name = '宋体'

#标题中文字体
run8.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run8.font.size = Pt(15)

#标题字体颜色
run8.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run8.bold = True



#定义表格熟悉(11行，3列)
t8 = doc1.add_table(rows=9, cols=2, style="Table Grid")

#定义表头
cells8 = t8.rows[0].cells
cells8[0].text='数据库性能信息'
cells8[1].text= '内容'

#设置表格宽度
cells8[1].width = Inches(10)


#第一行数据
t8.cell(1,0).text='会话信息'
t8.cell(1,1).text=f56.read()

#第二行数据
t8.cell(2,0).text='资源信息'
t8.cell(2,1).text=f57.read()

#第三行数据
t8.cell(3,0).text='顾问任务信息'
t8.cell(3,1).text=f58.read()

#第四行数据
t8.cell(4,0).text='SGA等信息'
t8.cell(4,1).text=f59.read()

#第五行数据
t8.cell(5,0).text='等待事件信息'
t8.cell(5,1).text=f60.read()

#第六行数据
t8.cell(6,0).text='内存信息'
t8.cell(6,1).text=f61.read()

#第七行数据
t8.cell(7,0).text='占内存进程信息'
t8.cell(7,1).text=f62.read()

#第八行数据
t8.cell(8,0).text='占CPU进程信息'
t8.cell(8,1).text=f63.read()
########################################



###########七:数据库参数信息############
#数据库巡检内容
paragraph9 = doc1.add_heading(level = 1)
run9 = paragraph9.add_run('七：数据库参数信息')

#标题英文字体
run9.font.name = '宋体'

#标题中文字体
run9.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run9.font.size = Pt(15)

#标题字体颜色
run9.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run9.bold = True


#定义表格熟悉(11行，3列)
t9 = doc1.add_table(rows=2, cols=2, style="Table Grid")

#定义表头
cells9 = t9.rows[0].cells
cells9[0].text='数据库参数信息'
cells9[1].text= '内容'

#设置表格宽度
cells9[1].width = Inches(9)


#第一行数据
t9.cell(1,0).text='需要优化的参数'
t9.cell(1,1).text=f64.read()
########################################




###########八:系统信息############
#数据库巡检内容
paragraph10 = doc1.add_heading(level = 1)
run10 = paragraph10.add_run('八：数据库系统信息')

#标题英文字体
run10.font.name = '宋体'

#标题中文字体
run10.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run10.font.size = Pt(15)

#标题字体颜色
run10.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run10.bold = True


#定义表格熟悉(11行，3列)
t10 = doc1.add_table(rows=4, cols=2, style="Table Grid")

#定义表头
cells10 = t10.rows[0].cells
cells10[0].text='数据库系统信息'
cells10[1].text= '内容'

#设置表格宽度
cells10[1].width = Inches(10)

#第1行数据
t10.cell(1,0).text='目录信息'
t10.cell(1,1).text=f65.read()
#第2行数据
t10.cell(2,0).text='挂载信息'
t10.cell(2,1).text=f66.read()
#第3行数据
t10.cell(3,0).text='自动挂载配置'
t10.cell(3,1).text=f67.read()
########################################


###########九:数据库日志信息############
#数据库巡检内容
paragraph11 = doc1.add_heading(level = 1)
run11 = paragraph11.add_run('九：数据库日志信息')

#标题英文字体
run11.font.name = '宋体'

#标题中文字体
run11.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

#字号：默认11号，改成15号
run11.font.size = Pt(15)

#标题字体颜色
run11.font.color.rgb = RGBColor(0, 0, 0)      #黑色

#字体加粗
run11.bold = True


#定义表格熟悉(11行，3列)
t11 = doc1.add_table(rows=5, cols=2, style="Table Grid")

#定义表头
cells11 = t11.rows[0].cells
cells11[0].text='数据库日志信息'
cells11[1].text= '内容'

#设置表格宽度
cells11[1].width = Inches(10)


#第一行数据
t11.cell(1,0).text='数据库告警日志'
t11.cell(1,1).text=f68.read()
#第二行数据
t11.cell(2,0).text='数据库集群日志'
t11.cell(2,1).text='不适用'
#第三行数据
t11.cell(3,0).text='数据库ASM日志'
t11.cell(3,1).text='不适用'
#第四行数据
t11.cell(4,0).text='操作系统日志'
t11.cell(4,1).text=f69.read()
########################################



#保存文档
doc1.save('20220620-AutoDBCheck系统数据库巡检报告-3.9.docx')



#关闭文件
f1.close()
f2.close()
f02.close()
f3.close()
f4.close()
f5.close()
f6.close()
f7.close()
f8.close()
f9.close()
f10.close()
f11.close()
f12.close()
f13.close()
f14.close()
f15.close()
f16.close()
f17.close()
f18.close()
f19.close()
f20.close()
f21.close()
f22.close()
f23.close()
f24.close()
f25.close()
f26.close()
f27.close()
f28.close()
f29.close()
f30.close()
f31.close()
f32.close()
f33.close()
f34.close()
f35.close()
f36.close()
f37.close()
f38.close()
f39.close()
f40.close()
f41.close()
f42.close()
f43.close()
f44.close()
f45.close()
f46.close()
f47.close()
f48.close()
f49.close()
f50.close()
f51.close()
f52.close()
f53.close()
f54.close()
f55.close()
f56.close()
f57.close()
f58.close()
f59.close()
f60.close()
f61.close()
f62.close()
f63.close()
f64.close()
f65.close()
f66.close()
f67.close()
f68.close()
f69.close()

print ("巡检报告已生成到D:\python\py目录下")



